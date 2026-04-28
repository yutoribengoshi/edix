#!/usr/bin/env python3
"""
Markdown → docx 出力モジュール

Edix のプレビュー用Markdownを Word (.docx) に変換。
A4・游明朝・25mmマージン・1.7行間・ページ番号付きで出力。

依存: python-docx, markdown, pymdown-extensions
"""
from __future__ import annotations

import re
import io
import xml.etree.ElementTree as ET
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================
# Markdown 拡張
# ============================================================
MD_EXTENSIONS = [
    'extra',
    'toc',
    'sane_lists',
    'pymdownx.tilde',
    'pymdownx.tasklist',
]

# ============================================================
# Word スタイル設定値
# ============================================================
FONT_BODY = '游明朝'
FONT_MONO = 'Consolas'
FONT_MONO_EAST = 'ＭＳ ゴシック'
BODY_FONT_SIZE = 11
LINE_SPACING = 1.7
PAGE_MARGIN_MM = 25
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297

HEADING_SIZES = {1: 18, 2: 14, 3: 12, 4: 11, 5: 11, 6: 11}
HEADING_BOLD = {1: True, 2: True, 3: True, 4: True, 5: False, 6: False}


# ============================================================
# Markdown 前処理
# ============================================================
def preprocess_markdown(md_text: str) -> str:
    """強制改ページ・添付資料連番"""
    md_text = re.sub(
        r'^\s*(?:<!--\s*page-break\s*-->|---page---)\s*$',
        '<div class="page-break"></div>',
        md_text, flags=re.MULTILINE
    )
    counter = [0]
    def attach_replacer(m):
        counter[0] += 1
        n = counter[0]
        zenkaku = str.maketrans('0123456789', '０１２３４５６７８９')
        return f'添付資料{str(n).translate(zenkaku)}'
    md_text = re.sub(r'添付資料【】', attach_replacer, md_text)
    return md_text


# ============================================================
# python-docx ユーティリティ
# ============================================================
def _set_run_font(run, size_pt=None, bold=False, italic=False,
                  strike=False, monospace=False):
    name = FONT_MONO if monospace else FONT_BODY
    east = FONT_MONO_EAST if monospace else FONT_BODY
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), east)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if strike:
        run.font.strike = True


def _set_paragraph_spacing(p, line_spacing=LINE_SPACING,
                           space_before=0, space_after=0):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def _shade_cell(cell, fill_hex='EFEFEF'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _setup_document(doc: Document):
    """A4・余白25mm・本文フォント設定 + フッターにページ番号"""
    section = doc.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(PAGE_MARGIN_MM)
    section.bottom_margin = Mm(PAGE_MARGIN_MM)
    section.left_margin = Mm(PAGE_MARGIN_MM)
    section.right_margin = Mm(PAGE_MARGIN_MM)

    # フッターにページ番号
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_run = fp.add_run()
    _set_run_font(fld_run, size_pt=10)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    fld_run._element.append(fldChar1)
    fld_run._element.append(instrText)
    fld_run._element.append(fldChar2)

    # 本文の既定スタイル
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(BODY_FONT_SIZE)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_BODY)
    rFonts.set(qn('w:hAnsi'), FONT_BODY)
    rFonts.set(qn('w:eastAsia'), FONT_BODY)


def _add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ============================================================
# HTML → ElementTree（簡易整形を経由）
# ============================================================
def _html_to_tree(html: str) -> ET.Element:
    """markdownライブラリ出力のHTMLをElementTreeに変換。
    XML として読めるよう、自閉じタグや実体参照を整形する。"""
    # void要素を自閉じに
    void_tags = ['br', 'hr', 'img', 'input', 'meta', 'link']
    for tag in void_tags:
        html = re.sub(rf'<{tag}([^>]*?)(?<!/)>', rf'<{tag}\1 />', html, flags=re.IGNORECASE)
    # 名前付き実体参照を数値参照に（既知のいくつか）
    entities = {'&nbsp;': '&#160;', '&laquo;': '&#171;', '&raquo;': '&#187;',
                '&copy;': '&#169;', '&reg;': '&#174;', '&hellip;': '&#8230;',
                '&mdash;': '&#8212;', '&ndash;': '&#8211;',
                '&lsquo;': '&#8216;', '&rsquo;': '&#8217;',
                '&ldquo;': '&#8220;', '&rdquo;': '&#8221;'}
    for k, v in entities.items():
        html = html.replace(k, v)
    # ルートで囲む
    wrapped = f'<root>{html}</root>'
    try:
        return ET.fromstring(wrapped)
    except ET.ParseError:
        # 失敗時：最低限のフォールバック
        wrapped2 = re.sub(r'&(?!#?\w+;)', '&amp;', wrapped)
        return ET.fromstring(wrapped2)


# ============================================================
# 文書ビルダー
# ============================================================
class DocxBuilder:
    def __init__(self, doc: Document):
        self.doc = doc
        # スタック状態
        self.list_stack = []  # [('ul'|'ol', counter)]

    def render_root(self, root: ET.Element):
        for child in root:
            self._render_block(child)

    # ------------------------------------------------------------
    # ブロック要素
    # ------------------------------------------------------------
    def _render_block(self, el: ET.Element):
        tag = el.tag.lower()

        if tag == 'p':
            p = self.doc.add_paragraph()
            _set_paragraph_spacing(p)
            self._render_inline(el, p)
            return

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            p = self.doc.add_paragraph()
            _set_paragraph_spacing(p, space_before=8, space_after=4)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._render_inline(el, p,
                                size=HEADING_SIZES.get(level, BODY_FONT_SIZE),
                                bold=HEADING_BOLD.get(level, False))
            return

        if tag == 'div':
            cls = el.get('class', '')
            if cls == 'page-break':
                _add_page_break(self.doc)
                return
            # その他のdivはコンテナとして子要素を処理
            for child in el:
                self._render_block(child)
            return

        if tag == 'blockquote':
            for child in el:
                if child.tag.lower() in ('p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Mm(8)
                    _set_paragraph_spacing(p)
                    self._render_inline(child, p, italic=True)
                else:
                    self._render_block(child)
            return

        if tag == 'pre':
            # コードブロック：<pre><code>...</code></pre>
            code_el = el.find('code')
            text = (code_el.text if code_el is not None else el.text) or ''
            for line in text.rstrip('\n').split('\n'):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Mm(5)
                _set_paragraph_spacing(p, line_spacing=1.2)
                run = p.add_run(line if line else ' ')
                _set_run_font(run, size_pt=10, monospace=True)
            return

        if tag in ('ul', 'ol'):
            self._render_list(el, tag, depth=0)
            return

        if tag == 'hr':
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('—' * 30)
            _set_run_font(run, size_pt=BODY_FONT_SIZE)
            return

        if tag == 'table':
            self._render_table(el)
            return

        if tag == 'dl':
            # 定義リスト：dt 太字、dd インデント
            for child in el:
                t = child.tag.lower()
                if t == 'dt':
                    p = self.doc.add_paragraph()
                    _set_paragraph_spacing(p)
                    self._render_inline(child, p, bold=True)
                elif t == 'dd':
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Mm(8)
                    _set_paragraph_spacing(p)
                    self._render_inline(child, p)
            return

        # その他：子要素を再帰
        for child in el:
            self._render_block(child)

    def _render_list(self, el: ET.Element, list_type: str, depth: int):
        counter = 1
        start_attr = el.get('start')
        if list_type == 'ol' and start_attr:
            try:
                counter = int(start_attr)
            except ValueError:
                pass
        for li in el:
            if li.tag.lower() != 'li':
                continue
            # マーカー作成
            indent_mm = 7 + depth * 7
            if list_type == 'ul':
                marker = '・'
            else:
                marker = f'{counter}. '
                counter += 1

            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(indent_mm)
            _set_paragraph_spacing(p)
            run = p.add_run(marker)
            _set_run_font(run, size_pt=BODY_FONT_SIZE)

            # li の中身：直接テキスト + インライン要素 + ネストされたリストやブロック
            self._render_inline(li, p, skip_block_children=True)

            # ネストされたリスト・ブロック要素を別段落で処理
            for child in li:
                ct = child.tag.lower()
                if ct in ('ul', 'ol'):
                    self._render_list(child, ct, depth + 1)
                elif ct in ('p', 'pre', 'blockquote', 'table'):
                    self._render_block(child)

    def _render_table(self, table_el: ET.Element):
        # 行を集める（thead と tbody を統合）
        rows = []  # list of (is_header, [cell_elements])
        def collect(parent, is_header_section):
            for child in parent:
                if child.tag.lower() == 'tr':
                    cells = [c for c in child if c.tag.lower() in ('th', 'td')]
                    has_header_cells = any(c.tag.lower() == 'th' for c in cells)
                    rows.append((is_header_section or has_header_cells, cells))
                elif child.tag.lower() in ('thead', 'tbody', 'tfoot'):
                    collect(child, child.tag.lower() == 'thead')

        collect(table_el, False)
        if not rows:
            return

        # 列数 = 最大セル数
        n_cols = max(len(cells) for _, cells in rows)
        n_rows = len(rows)

        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, (is_header, cells) in enumerate(rows):
            for c_idx, cell_el in enumerate(cells):
                cell = table.rows[r_idx].cells[c_idx]
                # セル内の既存の空段落をクリア
                cell.text = ''
                p = cell.paragraphs[0]
                _set_paragraph_spacing(p)
                self._render_inline(cell_el, p, bold=is_header)
                if is_header:
                    _shade_cell(cell)

    # ------------------------------------------------------------
    # インライン要素
    # ------------------------------------------------------------
    def _render_inline(self, el: ET.Element, paragraph,
                       size=BODY_FONT_SIZE, bold=False, italic=False,
                       strike=False, monospace=False, skip_block_children=False):
        """el の中身を paragraph に書き込む。
        skip_block_children=True なら、ブロック子要素（ul/ol/p等）はスキップ。"""
        # 直接テキスト
        if el.text:
            run = paragraph.add_run(el.text)
            _set_run_font(run, size_pt=size, bold=bold, italic=italic,
                          strike=strike, monospace=monospace)

        for child in el:
            tag = child.tag.lower()

            # スキップ対象（ブロック）
            if skip_block_children and tag in ('ul', 'ol', 'p', 'pre', 'blockquote', 'table', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                # 子はあとで別段落で処理されるので tail だけ拾う
                if child.tail:
                    run = paragraph.add_run(child.tail)
                    _set_run_font(run, size_pt=size, bold=bold, italic=italic,
                                  strike=strike, monospace=monospace)
                continue

            # 強調
            if tag in ('strong', 'b'):
                self._render_inline(child, paragraph, size=size, bold=True,
                                    italic=italic, strike=strike, monospace=monospace)
            elif tag in ('em', 'i'):
                self._render_inline(child, paragraph, size=size, bold=bold,
                                    italic=True, strike=strike, monospace=monospace)
            elif tag == 'code':
                self._render_inline(child, paragraph, size=size, bold=bold,
                                    italic=italic, strike=strike, monospace=True)
            elif tag in ('del', 's'):
                self._render_inline(child, paragraph, size=size, bold=bold,
                                    italic=italic, strike=True, monospace=monospace)
            elif tag == 'br':
                paragraph.add_run().add_break()
            elif tag == 'a':
                # リンク：テキストはそのまま、URLは括弧で添える（簡易）
                self._render_inline(child, paragraph, size=size, bold=bold,
                                    italic=italic, strike=strike, monospace=monospace)
            else:
                # その他は中身だけ取り出す
                self._render_inline(child, paragraph, size=size, bold=bold,
                                    italic=italic, strike=strike, monospace=monospace)

            # tail テキスト（タグの後ろのテキスト）
            if child.tail:
                run = paragraph.add_run(child.tail)
                _set_run_font(run, size_pt=size, bold=bold, italic=italic,
                              strike=strike, monospace=monospace)


# ============================================================
# メインエントリ
# ============================================================
def markdown_to_docx(md_text: str, output_path: Path = None) -> bytes:
    """Markdown を docx に変換。output_path 指定で保存も。返り値は docx の bytes。"""
    md_text = preprocess_markdown(md_text)
    md = markdown.Markdown(
        extensions=MD_EXTENSIONS,
        extension_configs={'toc': {'permalink': False, 'baselevel': 1}}
    )
    html = md.convert(md_text)
    root = _html_to_tree(html)

    doc = Document()
    _setup_document(doc)
    builder = DocxBuilder(doc)
    builder.render_root(root)

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if output_path:
        output_path.write_bytes(data)
    return data


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("Usage: docx_export.py <input.md> [output.docx]")
        sys.exit(1)
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2]) if len(sys.argv) >= 3 else src.with_suffix('.docx')
    md_text = src.read_text(encoding='utf-8')
    markdown_to_docx(md_text, dst)
    print(f"Wrote: {dst}")
